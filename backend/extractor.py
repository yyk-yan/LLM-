"""
extractor.py
文档内容提取模块 - 支持 .docx / .xlsx / .md / .txt,将整个文件内容提取为纯文本(用于后续 LLM 处理)
对于 Excel 文件，提供按关键词过滤行并返回结构化数据（表头 + 行字典）的能力
提供获取模板结构信息的功能（例如 Excel 的表头、Word 的段落和表格表头），用于后续填充
"""
import re
from pathlib import Path
"""读写xslx文件"""
import openpyxl
"""读写word文件"""
from docx import Document
"""优化一PDF支持模块"""
import pdfplumber
"""PyMuPDF"""
import fitz

import os
os.environ["FLAGS_use_mkldnn"] = "0"

"""
从word文件中提取所有文本,包括段落和表格内容
返回一个字符串，段落和表格行之间用换行分隔，表格内单元格用 | 连接
"""
def extract_text_from_docx(file_path: str) -> str:
    doc = Document(file_path)
    # 存储提取的文本片段
    parts = []
    # 遍历所有段落,把清理干净的文本，添加到 parts
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text.strip())
    # 遍历所有表格
    for table in doc.tables:
        for row in table.rows:
            # 提取一行中每个单元格的文本，并去除首尾空格
            row_data = [cell.text.strip() for cell in row.cells]
            if any(row_data):
                parts.append(" | ".join(row_data))
    # 将所有片段用换行符连接
    return "\n".join(parts)

"""
从Excel文件中提取所有工作表的内容为纯文本
每个工作表以[Sheet: 名称]开头，每行数据用|连接
"""
def extract_text_from_xlsx(file_path: str) -> str:
    # data_only=True 表示获取公式的计算值而非公式本身
    wb = openpyxl.load_workbook(file_path, data_only=True)
    parts = []
    # 遍历所有工作表
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        parts.append(f"[Sheet: {sheet_name}]")
        # 逐行读取单元格的值，转换成字符串
        for row in ws.iter_rows(values_only=True):
            row_data = [str(v) if v is not None else "" for v in row]
            if any(v.strip() for v in row_data):
                parts.append(" | ".join(row_data))
    return "\n".join(parts)

"""读取Markdown文件，返回原始文本内容"""
def extract_text_from_md(file_path: str) -> str:
    # 以 UTF-8 编码打开，忽略无法解码的字符
    with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
        return f.read()

"""读取txt文件，返回原始文本内容"""
def extract_text_from_txt(file_path: str) -> str:
    with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
        return f.read()


def extract_text_from_pdf(file_path: str) -> str:
    """
    智能提取PDF文本和表格：
    - 有文字层：pdfplumber提取文字+表格
    - 扫描件：EasyOCR识别文字
    """
    result = extract_pdf_with_tables(file_path)
    text = result["text"]
    tables = result["tables"]

    # 判断是否为扫描件
    if len(text.strip()) < 50:
        print("  检测到扫描件PDF，切换为OCR识别...")
        return extract_scanned_pdf(file_path)

    # 把表格转成文字追加到文本后面
    if tables:
        table_text = pdf_tables_to_text(tables)
        text = text + "\n\n[PDF表格内容]\n" + table_text

    return text


def extract_pdf_with_tables(file_path: str) -> dict:
    """
    从PDF中分别提取文字内容和表格内容
    返回格式：
    {
        "text": "页面文字...",
        "tables": [
            [["表头1", "表头2"], ["数据1", "数据2"]],  # 每个表格是二维list
            ...
        ]
    }
    """
    text_parts = []
    all_tables = []

    with pdfplumber.open(file_path) as pdf:
        for i, page in enumerate(pdf.pages):
            # 提取文字
            text = page.extract_text()
            if text and text.strip():
                text_parts.append(f"[第{i + 1}页]\n{text.strip()}")

            # 提取表格（返回二维list，每个元素是一行）
            tables = page.extract_tables()
            for table in tables:
                if table:  # 过滤空表格
                    all_tables.append(table)

    return {
        "text": "\n".join(text_parts),
        "tables": all_tables
    }


def pdf_tables_to_text(tables: list) -> str:
    """
    把表格二维list转成文字格式，方便送给LLM
    格式：表头1 | 表头2 | 表头3
          数据1 | 数据2 | 数据3
    """
    parts = []
    for table in tables:
        rows = []
        for row in table:
            # 过滤None，转字符串
            row_data = [str(cell) if cell else "" for cell in row]
            rows.append(" | ".join(row_data))
        parts.append("\n".join(rows))
    return "\n\n".join(parts)


def extract_scanned_pdf(file_path: str) -> str:
    """
    从扫描件PDF中提取文字，使用EasyOCR进行识别
    适用于无文字层的扫描PDF
    流程：PDF每页→渲染成图片→OCR识别→拼接文字
    """
    import easyocr
    import numpy as np
    # ch_sim=简体中文，en=英文，第一次运行会自动下载模型
    reader = easyocr.Reader(['ch_sim', 'en'], gpu=False)

    doc = fitz.open(file_path)
    parts = []

    for i, page in enumerate(doc):
        pix = page.get_pixmap(dpi=200)
        img_array = np.frombuffer(pix.samples, dtype=np.uint8).reshape(pix.height, pix.width, pix.n)

        # easyocr直接返回 [(bbox, text, confidence), ...]
        result = reader.readtext(img_array)

        if result:
            lines = [item[1] for item in result if item[1].strip()]
            if lines:
                parts.append(f"[第{i+1}页]\n" + "\n".join(lines))

    doc.close()
    return "\n".join(parts)

"""从Excel文件中按关键词条件过滤行，返回(headers, rows)"""
def extract_xlsx_rows(file_path: str, match_all: list = None, match_any: list = None) -> tuple:
    """按条件过滤xlsx，返回 (headers, rows) 结构化数据。
    match_all: 行必须包含所有关键词（AND）
    match_any: 行包含任意关键词即可（OR）
    两者都有时：match_all AND match_any 都满足"""
    wb = openpyxl.load_workbook(file_path, data_only=True)
    all_headers = []#and逻辑的表头
    all_rows = []#存储and逻辑匹配行
    #遍历所有sheet
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        headers = None
        for i, row in enumerate(ws.iter_rows(values_only=True)):
            row_data = [str(v) if v is not None else "" for v in row]
            #第一行作为表头（最终返回的 all_headers 只是第一个 sheet 的表头）
            if i == 0:
                headers = row_data
                if not all_headers:
                    all_headers = headers
                continue
            # 将整行所有单元格内容拼接成一个字符串，空格分隔
            row_str = " ".join(row_data)
            ok = True
            #如果 match_all 存在，要求所有关键词都出现在 row_str 中
            if match_all:
                ok = ok and all(kw in row_str for kw in match_all)
            #如果 match_any 存在，要求至少一个关键词出现
            #两者同时存在时，and 连接，即行必须同时满足两组条件
            if match_any:
                ok = ok and any(kw in row_str for kw in match_any)
            #存储匹配行到all_rows
            if ok:
                all_rows.append(dict(zip(headers, row_data)))
    print(f"  xlsx结构化过滤: 匹配{len(all_rows)}行")
    #返回第一个 sheet 的表头，以及所有匹配行的字典列表
    return all_headers, all_rows

"""从表格型PDF中结构化提取，返回(headers, rows)，用法与extract_xlsx_rows一致"""
def extract_pdf_rows(file_path: str, match_all: list = None, match_any: list = None) -> tuple:
    """
    用 pdfplumber 的 extract_tables 按表格读取，保留空单元格的列位置；
    自动识别表头、跳过跨页重复表头与注释行。识别不到表格时返回 ([], [])，
    由 agent 回退到文本提取路径。
    """
    headers = None
    all_rows = []
    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            for table in page.extract_tables():
                if not table:
                    continue
                for raw in table:
                    cells = [(str(c).strip() if c is not None else "") for c in raw]
                    # 几乎全空的行（如页脚注释整段挤进一个单元格）跳过
                    if sum(1 for c in cells if c) < 2:
                        continue
                    # 第一行非空作为表头
                    if headers is None:
                        headers = cells
                        continue
                    # 跨页重复出现的表头行跳过
                    if cells == headers:
                        continue
                    # 列数与表头对不上的行（说明、残行）跳过
                    if len(cells) != len(headers):
                        continue
                    row_str = " ".join(cells)
                    ok = True
                    if match_all:
                        ok = ok and all(kw in row_str for kw in match_all)
                    if match_any:
                        ok = ok and any(kw in row_str for kw in match_any)
                    if ok:
                        all_rows.append(dict(zip(headers, cells)))
    print(f"  PDF结构化表格提取: 匹配{len(all_rows)}行")
    return (headers or []), all_rows

"""根据文件扩展名自动选择对应的提取函数"""
def extract_file(file_path: str) -> str:
    path = Path(file_path)
    ext = path.suffix.lower()
    if ext == ".docx":
        return extract_text_from_docx(file_path)
    elif ext in (".xlsx", ".xls"):
        return extract_text_from_xlsx(file_path)
    elif ext == ".md":
        return extract_text_from_md(file_path)
    elif ext == ".txt":
        return extract_text_from_txt(file_path)
    elif ext == ".pdf":
        return extract_text_from_pdf(file_path)
    else:
        raise ValueError(f"不支持的文件类型: {ext}")


"""
获取xlsx和word模板的结构信息
返回字典，描述有哪些列需要填
"""
def get_xlsx_structure(file_path: str) -> dict:
    """获取xlsx模板的结构信息（sheet名、表头、行数）"""
    wb = openpyxl.load_workbook(file_path, data_only=True)
    structure = {}
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        headers = []
        rows = []
        for i, row in enumerate(ws.iter_rows(values_only=True)):
            row_data = [str(v) if v is not None else "" for v in row]
            if i == 0:
                headers = row_data
            else:
                if any(v.strip() for v in row_data):
                    rows.append(row_data)
        structure[sheet_name] = {"headers": headers, "existing_rows": rows}
    return structure

"""获取 Word 模板的结构信息,{"paragraphs":[段落1,段落2,...],"tables":[{"index":0,"headers":[表头1,表头2,...],"row_counts":5}]}"""
def get_docx_structure(file_path: str) -> dict:
    """获取docx模板的结构信息（段落描述、表格表头）"""
    doc = Document(file_path)
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    tables = []
    for i, table in enumerate(doc.tables):
        if table.rows:
            headers = [cell.text.strip() for cell in table.rows[0].cells]
            tables.append({"index": i, "headers": headers, "row_count": len(table.rows)})
    return {"paragraphs": paragraphs, "tables": tables}
