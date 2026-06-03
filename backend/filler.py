"""
filler.py
表格填写模块
将agent提取的结构化数据写入xlsx或docx模板
写入时仅修改单元格值，完整保留原模板格式属性；
支持精确匹配与去空格模糊匹配两种字段对齐策略，docx模板不足时动态复制行样式扩展表格。
"""
import copy
import shutil
from pathlib import Path
import openpyxl
from docx import Document

"""
写入xlsx模板，返回输出文件的路径
1.复制模板文件到输出路径（避免修改原模板）。
2.加载输出文件的工作簿。
3.对fill_data中的每个工作表：
    定位工作表（如果sheet名不存在，则使用当前活动sheet）。
    读取第一行作为表头（列名）。
    从第二行开始查找第一个全空的行作为起始写入行。
    逐行写入数据，支持精确匹配和去空格模糊匹配。
4.保存工作簿。
"""
def fill_xlsx(template_path: str, output_path: str, fill_data: dict) -> str:
    """
    fill_data 格式:
    {
      "Sheet1": [
        {"col1": "val1", "col2": "val2"},
        ...
      ]
    }
    """
    # 复制模板到输出路径（保留原模板不变）
    shutil.copy2(template_path, output_path)
     # 加载输出文件的工作簿
    wb = openpyxl.load_workbook(output_path)

    # 遍历 fill_data 中的每个sheet
    for sheet_name, rows in fill_data.items():
         # 如果工作表名不存在于工作簿中，则使用当前活动工作表
        if sheet_name not in wb.sheetnames:
            # 尝试用第一个sheet
            ws = wb.active
        else:
            ws = wb[sheet_name]

        # 找表头行（第一行）
        headers = [str(ws.cell(1, c).value or "").strip() for c in range(1, ws.max_column + 2)]

        # 找第一个空行
        start_row = 2
        for r in range(2, ws.max_row + 2):
            # 如果该行所有单元格（从第1列到表头长度列）都为空，则找到空行
            if all(ws.cell(r, c).value is None for c in range(1, len(headers) + 1)):
                start_row = r
                break

         # 逐行写入数据
        for row_idx, row_data in enumerate(rows):
            r = start_row + row_idx
            for col_idx, header in enumerate(headers):
                if not header:
                    continue
                # 精确匹配（优先）或模糊匹配
                value = row_data.get(header)
                if value is None:
                    # 精确匹配失败，尝试去空格匹配
                    for k, v in row_data.items():
                        if k.replace(" ", "") == header.replace(" ", ""):
                            value = v
                            break
                if value is not None:
                    ws.cell(r, col_idx + 1, value)
    # 保存工作簿
    wb.save(output_path)
    return output_path

"""
写入word模板，返回输出文件的路径
1.复制模板文件到输出路径（避免修改原模板）。
2.加载输出文件的document对象。
3.对fill_data中的每个表格：
    从key中提取表格索引（例如 "table_0" -> 0）。
    获取该表格对象，读取第一行作为表头。
    查找表格中第一个全空的行（从第二行开始）作为起始写入行。
    逐行写入数据，如果表格行数不够，调用_add_row动态添加新行（复制最后一行样式并清空内容）。
4.保存文档。
"""
def fill_docx(template_path: str, output_path: str, fill_data: dict) -> str:
    """
    fill_data 格式:
    {
      "table_0": [
        {"col1": "val1", "col2": "val2"},
        ...
      ],
      "table_1": [...]
    }
    """
    shutil.copy2(template_path, output_path)
    doc = Document(output_path)

    for key, rows in fill_data.items():
        # key: "table_0", "table_1", ...
        try:
            table_idx = int(key.split("_")[-1])
        except Exception:
            continue

        if table_idx >= len(doc.tables):
            continue

        table = doc.tables[table_idx]
        headers = [cell.text.strip() for cell in table.rows[0].cells]

        # 找第一个空行
        start_row = 1
        for r_idx, row in enumerate(table.rows[1:], start=1):
            if all(not cell.text.strip() for cell in row.cells):
                start_row = r_idx
                break

        for row_idx, row_data in enumerate(rows):
            r_idx = start_row + row_idx
            # 如果行不够，复制最后一行样式添加新行
            while r_idx >= len(table.rows):
                _add_row(table)

            row = table.rows[r_idx]
            for col_idx, header in enumerate(headers):
                if not header or col_idx >= len(row.cells):
                    continue
                value = row_data.get(header)
                if value is None:
                    for k, v in row_data.items():
                        if k.replace(" ", "") == header.replace(" ", ""):
                            value = v
                            break
                if value is not None:
                    row.cells[col_idx].text = str(value)

    doc.save(output_path)
    return output_path

"""
动态添加行（底层XML操作）
Word文件本质是XML
python-docx提供的高层API不支持添加行，这里直接操作底层XML
深拷贝最后一行（连同它的字体、边框等格式），然后把所有文字清空，再追加到表格末尾
这就是格式无损数据写入的实现原理
"""
def _add_row(table):
    """在表格末尾添加一行（复制最后一行的XML结构）"""
    from docx.oxml.ns import qn
    import copy
    last_row = table.rows[-1]._tr
    new_row = copy.deepcopy(last_row)
    # 清空单元格内容
    for tc in new_row.findall(qn("w:tc")):
        for p in tc.findall(qn("w:p")):
            for r in p.findall(qn("w:r")):
                for t in r.findall(qn("w:t")):
                    t.text = ""
    table._tbl.append(new_row)
